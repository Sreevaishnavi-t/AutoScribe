import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm
from google import genai
import io
import tempfile
import os
import PyPDF2
from pdf2docx import Converter
from fpdf import FPDF

st.set_page_config(page_title="Doc Maker", layout="wide")

# CSS for Transitions and styling
st.markdown("""
<style>
    /* Fade-in animation for main container */
    .stApp > header {
        background-color: transparent;
    }
    .stApp {
        animation: fadeIn 0.8s ease-in-out;
    }
    @keyframes fadeIn {
        0% { opacity: 0; transform: translateY(10px); }
        100% { opacity: 1; transform: translateY(0); }
    }
    /* Button hover transitions */
    .stButton>button {
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: scale(1.02);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    /* Expander animation */
    .streamlit-expanderHeader {
        transition: background-color 0.3s ease;
    }
</style>
""", unsafe_allow_html=True)

st.title("📄 AutoScribe")

# --- Templates ---
TEMPLATES = {
    "Custom": "",
    "Resume Formatter": "1. Contact Information\n2. Professional Summary\n3. Work Experience\n4. Education\n5. Skills",
    "Meeting Notes": "1. Attendees\n2. Agenda Items\n3. Discussion\n4. Action Items\n5. Next Meeting",
    "Assignment Formatter": "1. Title & Student Info\n2. Abstract\n3. Introduction\n4. Methodology\n5. Results\n6. Discussion\n7. References",
    "Research Notes": "1. Source Reference\n2. Key Themes\n3. Important Quotes\n4. Personal Analysis\n5. Further Questions",
    "Blog Structure": "1. Catchy Title\n2. Hook/Introduction\n3. Main Point 1\n4. Main Point 2\n5. Main Point 3\n6. Conclusion & Call to Action"
}

# --- Utilities ---
def extract_text_from_file(uploaded_file):
    if not uploaded_file:
        return ""
    if uploaded_file.name.endswith('.txt'):
        return uploaded_file.getvalue().decode('utf-8', errors='ignore')
    elif uploaded_file.name.endswith('.docx'):
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif uploaded_file.name.endswith('.pdf'):
        try:
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {e}")
            return ""
    return ""

def generate_content(title, outline, api_key):
    if not api_key:
        return f"# {title}\n\n## Introduction\n\nThis is generated dummy text because no API key was provided. In a real scenario, the AI would generate a full introduction about {title} based on your outline.\n\n## Body Section\n\nLorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.\n\n## Conclusion\n\nThis concludes the mock document. If you want real text, paste a free Gemini API key in the sidebar."
    try:
        client = genai.Client(api_key=api_key)
        prompt = f"Write a comprehensive document titled '{title}'.\n\nHere is the outline to follow:\n{outline}\n\nWrite the content section by section. Use Markdown formatting: ## for Headings and standard text for paragraphs. Do not include a title at the very top as it will be added separately."
        response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating content: {e}")
        return None

def create_document(title, content_md, template_file, apply_overrides, font_family, body_size, h1_size, selected_color):
    if template_file:
        doc = Document(template_file)
        doc.add_page_break()
    else:
        doc = Document()
        # Apply standard margins if no template
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
            
    title_para = doc.add_paragraph(title, style='Title')
    title_para.alignment = 1 
    
    lines = content_md.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('## '):
            p = doc.add_paragraph(line.replace('## ', ''), style='Heading 1')
            if apply_overrides:
                run = p.runs[0] if p.runs else p.add_run(p.text)
                if p.runs and len(p.runs) > 1:
                   run.text = p.text
                   for r in p.runs[1:]: r.text = ''
                run.font.name = font_family
                run.font.size = Pt(h1_size)
                run.font.color.rgb = RGBColor(*selected_color)
        elif line.startswith('### '):
            p = doc.add_paragraph(line.replace('### ', ''), style='Heading 2')
            if apply_overrides:
                run = p.runs[0] if p.runs else p.add_run(p.text)
                if p.runs and len(p.runs) > 1:
                   run.text = p.text
                   for r in p.runs[1:]: r.text = ''
                run.font.name = font_family
                run.font.size = Pt(max(12, h1_size - 2))
                run.font.color.rgb = RGBColor(*selected_color)
        elif line.startswith('# '):
            p = doc.add_paragraph(line.replace('# ', ''), style='Heading 1')
            if apply_overrides:
                run = p.runs[0] if p.runs else p.add_run(p.text)
                if p.runs and len(p.runs) > 1:
                   run.text = p.text
                   for r in p.runs[1:]: r.text = ''
                run.font.name = font_family
                run.font.size = Pt(h1_size)
                run.font.color.rgb = RGBColor(*selected_color)
        else:
            p = doc.add_paragraph(line, style='Normal')
            if apply_overrides:
                if p.runs:
                    for run in p.runs:
                        run.font.name = font_family
                        run.font.size = Pt(body_size)
                else:
                    run = p.add_run(line)
                    p.text = ""
                    run.font.name = font_family
                    run.font.size = Pt(body_size)
    return doc

def clean_document(uploaded_messy_file, font_family, body_size):
    clean_doc = Document()
    
    # 1. Enforce Print-Ready Canvas (A4 & Margins)
    for section in clean_doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Extract text/paragraphs depending on file type
    if uploaded_messy_file.name.endswith('.docx'):
        doc = Document(uploaded_messy_file)
        paragraphs = [p.text for p in doc.paragraphs]
        styles = [p.style.name for p in doc.paragraphs]
        alignments = [p.alignment for p in doc.paragraphs]
    else:
        text = extract_text_from_file(uploaded_messy_file)
        paragraphs = text.split('\n')
        styles = ['Normal'] * len(paragraphs)
        alignments = [None] * len(paragraphs)

    # 2. Clean Paragraphs
    blank_count = 0
    for i, text in enumerate(paragraphs):
        text = text.strip()
        if not text:
            blank_count += 1
            if blank_count > 1:
                continue # Skip excessive blank lines
        else:
            blank_count = 0
            
        try:
            new_p = clean_doc.add_paragraph(text, style=styles[i])
        except KeyError:
            new_p = clean_doc.add_paragraph(text, style='Normal')
            
        if alignments[i] is not None:
            new_p.alignment = alignments[i]
        
        # Normalize font for body text
        if new_p.style.name == 'Normal':
            if new_p.runs:
                for run in new_p.runs:
                    run.font.name = font_family
                    run.font.size = Pt(body_size)
            else:
                run = new_p.add_run(text)
                new_p.text = ""
                run.font.name = font_family
                run.font.size = Pt(body_size)
                
    return clean_doc

# --- UI Toggles ---
mode = st.radio("Select Mode", ["Generate New Document", "Clean Existing Document", "File Converter"], horizontal=True)
st.divider()

# Sidebar for Settings
with st.sidebar:
    st.header("⚙️ Settings")
    api_key = st.text_input("Enter your gemini API key", type="password", help="Get a free key from aistudio.google.com")
    st.divider()
    
    # Initialize defaults to avoid UnboundLocalError
    font_family = "Arial"
    body_size = 12
    h1_size = 18
    selected_color = (0,0,0)
    override_formatting = False
    
    if mode == "Generate New Document":
        st.subheader("Formatting Overrides")
        st.caption("Use this if your reference document doesn't have clean styles, or if you aren't uploading one.")
        override_formatting = st.checkbox("Override Reference Styles", value=False)

        if override_formatting:
            font_family = st.selectbox("Font Family", ["Arial", "Times New Roman", "Calibri", "Helvetica"])
            body_size = st.number_input("Body Text Size (pt)", min_value=8, max_value=24, value=12)
            h1_size = st.number_input("Heading 1 Size (pt)", min_value=12, max_value=36, value=18)
            color_choice = st.selectbox("Heading Color", ["Black", "Dark Blue", "Dark Red", "Dark Green"])
            color_map = {"Black": (0, 0, 0), "Dark Blue": (0, 0, 139), "Dark Red": (139, 0, 0), "Dark Green": (0, 100, 0)}
            selected_color = color_map[color_choice]
    elif mode == "Clean Existing Document":
        st.subheader("Cleaner Settings")
        st.caption("Settings for formatting the messy document.")
        font_family = st.selectbox("Normalize Body Font to:", ["Arial", "Times New Roman", "Calibri", "Helvetica"])
        body_size = st.number_input("Normalize Body Size to (pt):", min_value=8, max_value=24, value=12)

# --- Main App Logic ---
if mode == "Generate New Document":
    st.write("Upload a reference document, or use the settings panel to generate a perfectly formatted, print-ready document.")
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("Document Content")
        doc_title = st.text_input("Document Title", placeholder="e.g., The History of Artificial Intelligence")
        
        template_choice = st.selectbox("Choose a Template / Structure", list(TEMPLATES.keys()))
        
        outline_file = st.file_uploader("Or Upload an Outline (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"], help="Drag and drop an outline file to auto-populate.")
        
        # Determine the outline text based on file upload or template choice
        if outline_file:
            default_outline = extract_text_from_file(outline_file)
        else:
            default_outline = TEMPLATES[template_choice]
            
        doc_outline = st.text_area("Outline / Table of Contents", value=default_outline, height=200)
        
    with col2:
        st.subheader("Reference Template")
        uploaded_file = st.file_uploader("Upload Reference Document (.docx only)", type=["docx"], key="gen_up", help="Must be .docx to extract styles.")

    if st.button("Generate Document", type="primary", use_container_width=True):
        if not doc_title:
            st.warning("Please provide a Document Title.")
        else:
            with st.spinner("Generating document content..."):
                markdown_content = generate_content(doc_title, doc_outline, api_key)
            if markdown_content:
                with st.spinner("Applying formatting and creating Word document..."):
                    try:
                        doc = create_document(doc_title, markdown_content, uploaded_file, override_formatting, font_family, body_size, h1_size, selected_color)
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.success("Document generated successfully!")
                        
                        st.markdown("### Export Options")
                        # Show export options horizontally
                        ex1, ex2, ex3 = st.columns(3)
                        with ex1:
                            st.download_button(
                                label="📥 Download DOCX",
                                data=buffer,
                                file_name=f"{doc_title.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        with ex2:
                            pdf = FPDF()
                            pdf.add_page()
                            pdf.set_auto_page_break(auto=True, margin=15)
                            pdf.set_font("helvetica", size=12)
                            # Safe encoding for basic PDF
                            for line in markdown_content.split('\n'):
                                pdf.multi_cell(0, 8, text=line.encode('latin-1', 'replace').decode('latin-1'))
                            pdf_bytes = pdf.output()
                            st.download_button(
                                label="📥 Download PDF",
                                data=io.BytesIO(pdf_bytes),
                                file_name=f"{doc_title.replace(' ', '_')}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        with ex3:
                            st.download_button(
                                label="📥 Download Markdown",
                                data=markdown_content,
                                file_name=f"{doc_title.replace(' ', '_')}.md",
                                mime="text/markdown",
                                use_container_width=True
                            )
                        
                        st.markdown("---")
                        with st.expander("📋 View / Copy Formatted Text"):
                            st.code(markdown_content, language="markdown")
                            
                    except Exception as e:
                        st.error(f"Error creating document: {e}")

elif mode == "Clean Existing Document":
    st.write("Upload a messy document. The app will force A4 sizing, 1-inch margins, delete random blank pages, and normalize fonts for printing.")
    st.subheader("Upload Messy Document")
    messy_file = st.file_uploader("Upload Document (.docx, .txt, .pdf)", type=["docx", "txt", "pdf"], key="clean_up", help="Drag and drop your messy file here.")
    
    if st.button("Clean and Format Document", type="primary", use_container_width=True):
        if not messy_file:
            st.warning("Please upload a document to clean.")
        else:
            with st.spinner("Cleaning document layout..."):
                try:
                    cleaned_doc = clean_document(messy_file, font_family, body_size)
                    buffer = io.BytesIO()
                    cleaned_doc.save(buffer)
                    buffer.seek(0)
                    st.success("Document cleaned successfully! (A4 Size, 1-inch Margins, No blank pages)")
                    st.download_button(
                        label="📥 Download Cleaned .docx",
                        data=buffer,
                        file_name=f"Cleaned_{messy_file.name.split('.')[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                except Exception as e:
                    st.error(f"Error cleaning document: {e}")

elif mode == "File Converter":
    st.write("Convert between Word (.docx) and PDF formats. **Drag and drop** supported.")
    convert_mode = st.radio("Conversion Type", ["Word to PDF", "PDF to Word"], horizontal=True)
    
    if convert_mode == "Word to PDF":
        st.info("Note: To ensure server compatibility, this extracts text from your Word document and generates a basic PDF. Complex tables and images may be excluded.")
        convert_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"], key="w2p")
        if st.button("Convert to PDF", type="primary"):
            if convert_file:
                with st.spinner("Converting to PDF..."):
                    try:
                        text = extract_text_from_file(convert_file)
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_auto_page_break(auto=True, margin=15)
                        pdf.set_font("helvetica", size=12)
                        for line in text.split('\n'):
                            pdf.multi_cell(0, 8, text=line.encode('latin-1', 'replace').decode('latin-1'))
                        
                        pdf_bytes = pdf.output()
                        buffer = io.BytesIO(pdf_bytes)
                        
                        st.success("Converted successfully!")
                        st.download_button(
                            label="📥 Download PDF", 
                            data=buffer, 
                            file_name=f"{convert_file.name.split('.')[0]}.pdf", 
                            mime="application/pdf"
                        )
                    except Exception as e:
                        st.error(f"Error during conversion: {e}")
            else:
                st.warning("Please upload a file.")
    else:
        st.info("Converts PDF layout and text back into an editable Word document.")
        convert_file = st.file_uploader("Upload PDF Document (.pdf)", type=["pdf"], key="p2w")
        if st.button("Convert to Word", type="primary"):
            if convert_file:
                with st.spinner("Converting to Word (this may take a moment)..."):
                    try:
                        # pdf2docx needs an actual file path
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                            temp_pdf.write(convert_file.getvalue())
                            temp_pdf_path = temp_pdf.name
                            
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                            temp_docx_path = temp_docx.name
                            
                        cv = Converter(temp_pdf_path)
                        cv.convert(temp_docx_path)
                        cv.close()
                        
                        with open(temp_docx_path, "rb") as f:
                            docx_bytes = f.read()
                            
                        os.remove(temp_pdf_path)
                        os.remove(temp_docx_path)
                        
                        st.success("Converted successfully!")
                        st.download_button(
                            label="📥 Download DOCX", 
                            data=docx_bytes, 
                            file_name=f"{convert_file.name.split('.')[0]}.docx", 
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Error during conversion: {e}")
            else:
                st.warning("Please upload a file.")
